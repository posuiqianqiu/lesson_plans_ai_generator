from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml

# 创建新的Word文档
doc = Document()

# 设置文档字体
doc.styles['Normal'].font.name = u'宋体'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
doc.styles['Normal'].font.size = Pt(12)

# 添加标题
title = doc.add_paragraph()
title_run = title.add_run('{{course_name}} 教案')
title_run.font.size = Pt(18)
title_run.font.bold = True
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加空行
doc.add_paragraph()

# 创建统一的大表格（14行 x 4列）
main_table = doc.add_table(rows=14, cols=4)
main_table.alignment = WD_TABLE_ALIGNMENT.CENTER
main_table.autofit = True

# 设置每列宽度
for row in main_table.rows:
    for i, cell in enumerate(row.cells):
        if i == 0 or i == 1:
            cell.width = Inches(1.2)
        elif i == 2:
            cell.width = Inches(3)
        else:
            cell.width = Inches(1.2)

# 第1行：基本信息表头
basic_headers = ['周次', '课次', '章节内容', '课时']
for i, header in enumerate(basic_headers):
    cell = main_table.cell(0, i)
    cell.text = header
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cell.paragraphs[0].runs:
        run.font.bold = True
        run.font.size = Pt(12)

# 第2行：基本信息数据
basic_data = ['{{week}}', '{{lesson}}', '{{chapter_content}}', '{{class_hours}}']
for i, data in enumerate(basic_data):
    cell = main_table.cell(1, i)
    cell.text = data
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# 第3行：单元教学目标（跨4列）
# 合并单元格
main_table.cell(2, 0).merge(main_table.cell(2, 3))
obj_header = main_table.cell(2, 0)
obj_header.text = '单元教学目标'
obj_header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in obj_header.paragraphs[0].runs:
    run.font.bold = True
    run.font.size = Pt(14)

# 第4行：教学目标内容（跨4列）
main_table.cell(3, 0).merge(main_table.cell(3, 3))
obj_content = main_table.cell(3, 0)
obj_content.text = '{{单元教学目标}}'
obj_content.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

# 第5行：教学重点和难点（左右分2列）
# 教学重点（跨2列）
main_table.cell(4, 0).merge(main_table.cell(4, 1))
key_header = main_table.cell(4, 0)
key_header.text = '教学重点'
key_header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in key_header.paragraphs[0].runs:
    run.font.bold = True
    run.font.size = Pt(13)

# 教学难点（跨2列）
main_table.cell(4, 2).merge(main_table.cell(4, 3))
diff_header = main_table.cell(4, 2)
diff_header.text = '教学难点'
diff_header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in diff_header.paragraphs[0].runs:
    run.font.bold = True
    run.font.size = Pt(13)

# 第6行：重点难点内容
main_table.cell(5, 0).merge(main_table.cell(5, 1))
key_content = main_table.cell(5, 0)
key_content.text = '{{教学重点}}'
key_content.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

main_table.cell(5, 2).merge(main_table.cell(5, 3))
diff_content = main_table.cell(5, 2)
diff_content.text = '{{教学难点}}'
diff_content.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

# 第7行：教学活动（跨4列）
main_table.cell(6, 0).merge(main_table.cell(6, 3))
act_header = main_table.cell(6, 0)
act_header.text = '教学活动'
act_header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in act_header.paragraphs[0].runs:
    run.font.bold = True
    run.font.size = Pt(14)

# 第8行：教学活动内容（跨4列）
main_table.cell(7, 0).merge(main_table.cell(7, 3))
act_content = main_table.cell(7, 0)
act_content.text = '{{教学活动}}'
act_content.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

# 第9行：教学资源（跨4列）
main_table.cell(8, 0).merge(main_table.cell(8, 3))
res_header = main_table.cell(8, 0)
res_header.text = '教学资源'
res_header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in res_header.paragraphs[0].runs:
    run.font.bold = True
    run.font.size = Pt(14)

# 第10行：教学资源内容（跨4列）
main_table.cell(9, 0).merge(main_table.cell(9, 3))
res_content = main_table.cell(9, 0)
res_content.text = '{{教学资源}}'
res_content.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

# 第11行：教学反思（跨4列）
main_table.cell(10, 0).merge(main_table.cell(10, 3))
ref_header = main_table.cell(10, 0)
ref_header.text = '教学反思'
ref_header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in ref_header.paragraphs[0].runs:
    run.font.bold = True
    run.font.size = Pt(14)

# 第12行：教学反思内容（跨4列）
main_table.cell(11, 0).merge(main_table.cell(11, 3))
ref_content = main_table.cell(11, 0)
ref_content.text = '{{教学反思}}'
ref_content.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

# 第13行：教学评价（跨4列）
main_table.cell(12, 0).merge(main_table.cell(12, 3))
eval_header = main_table.cell(12, 0)
eval_header.text = '教学评价'
eval_header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in eval_header.paragraphs[0].runs:
    run.font.bold = True
    run.font.size = Pt(14)

# 第14行：教学评价内容（跨4列）
main_table.cell(13, 0).merge(main_table.cell(13, 3))
eval_content = main_table.cell(13, 0)
eval_content.text = '{{教学评价}}'
eval_content.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

# 设置表格边框
def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = parse_xml(r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/><w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/></w:tblBorders>')
    tblPr.append(tblBorders)

set_table_borders(main_table)

# 保存文档
doc.save('test_data/template_single_table.docx')
print("单张表格格式模板已创建成功！")