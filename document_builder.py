from docx import Document
import os
from tqdm import tqdm

class DocumentBuilder:
    """文档组装器：按周次和课次批量生成Word格式教案"""
    
    def __init__(self, template_path):
        """
        初始化文档组装器
        :param template_path: 教案模板文件路径
        """
        self.template_path = template_path

    def _replace_text_in_doc(self, doc, replacements):
        """在文档的段落和表格中执行文本替换"""
        for p in doc.paragraphs:
            for key, value in replacements.items():
                if key in p.text:
                    p.text = p.text.replace(key, value)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in p.text:
                                p.text = p.text.replace(key, value)
        
    def build_lesson_plan(self, lesson_data, ai_generated_content, output_path):
        """
        生成单个教案文档
        :param lesson_data: 课程数据
        :param ai_generated_content: AI生成的内容字典
        :param output_path: 输出文件路径
        """
        try:
            doc = Document(self.template_path)
            
            # 准备替换数据
            replacements = {
                "{{week}}": str(lesson_data['week']),
                "{{lesson}}": str(lesson_data['lesson']),
                "{{course_name}}": lesson_data['课程名称'],
                "{{chapter_content}}": lesson_data['章节内容'],
                "{{class_hours}}": str(lesson_data['课时'])
            }
            
            # 添加AI生成的内容
            for field, content in ai_generated_content.items():
                replacements[f"{{{{{field}}}}}"] = content
            
            # 执行文本替换
            self._replace_text_in_doc(doc, replacements)
            
            # 保存文档
            doc.save(output_path)
            
        except Exception as e:
            print(f"生成教案失败: {e}")
            raise e
    
    def build_batch_lesson_plans(self, schedule_data, ai_generator, syllabus_data=None):
        """
        批量生成教案
        :param schedule_data: 教学进度表数据
        :param ai_generator: AI生成器实例
        :param syllabus_data: 教学大纲数据（可选）
        """
        if not os.path.exists('lesson_plans'):
            os.makedirs('lesson_plans')
        
        for lesson_data in tqdm(schedule_data, desc="生成教案"):
            try:
                # 生成AI内容
                ai_content = {}
                for field in ["单元教学目标", "教学重点", "教学难点", "教学方法", "教学过程", "作业布置"]:
                    ai_content[field] = ai_generator.generate_content(field, lesson_data, syllabus_data)
                
                # 生成文档
                output_filename = f"第{lesson_data['week']}周第{lesson_data['lesson']}次课教案.docx"
                output_path = os.path.join('lesson_plans', output_filename)
                
                self.build_lesson_plan(lesson_data, ai_content, output_path)
                
            except Exception as e:
                print(f"生成第{lesson_data['week']}周第{lesson_data['lesson']}次课教案失败: {e}")
                continue