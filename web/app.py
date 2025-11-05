from fastapi import FastAPI, File, UploadFile, HTTPException, WebSocket, WebSocketDisconnect, Form
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from fastapi.requests import Request
from fastapi.responses import HTMLResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import os
import sys
import uuid
import json
import asyncio
import time
from typing import Dict, List
import aiofiles
from pathlib import Path

# 设置默认环境变量
os.environ.setdefault('OLLAMA_MODEL', 'qwen3:1.7b')
os.environ.setdefault('OLLAMA_HOST', 'http://localhost:11434')

# 获取项目根目录
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.append(PROJECT_ROOT)

# 确保工作目录正确
if not os.getcwd().endswith('教案ai生成器'):
    os.chdir(PROJECT_ROOT)

print(f"项目根目录: {PROJECT_ROOT}")
print(f"当前工作目录: {os.getcwd()}")

from data_parser import DataParser
from ai_generator import AIGenerator
from document_builder import DocumentBuilder

# Pydantic模型
class ParseRequest(BaseModel):
    file_id: str

class GenerateRequest(BaseModel):
    schedule_file_id: str
    syllabus_file_id: str = None
    template_file_id: str = None
    week_range: str = None

app = FastAPI(title="教案AI生成器", description="高职院校教案智能生成系统")

# 添加 CORS 中间件
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # 允许所有来源，生产环境中应该限制为特定域名
    allow_credentials=True,
    allow_methods=["*"],  # 允许所有 HTTP 方法
    allow_headers=["*"],  # 允许所有请求头
)

# 静态文件和模板配置
static_dir = os.path.join(os.path.dirname(__file__), "static")
templates_dir = os.path.join(os.path.dirname(__file__), "templates")
uploads_dir = os.path.join(os.path.dirname(__file__), "uploads")

# 确保使用绝对路径
uploads_dir = os.path.abspath(uploads_dir)

# 确保必要的目录存在
os.makedirs(static_dir, exist_ok=True)
os.makedirs(templates_dir, exist_ok=True)
os.makedirs(uploads_dir, exist_ok=True)

app.mount("/static", StaticFiles(directory=static_dir), name="static")
templates = Jinja2Templates(directory=templates_dir)

# 全局变量存储上传文件和任务状态
uploaded_files: Dict[str, Dict] = {}
generation_tasks: Dict[str, Dict] = {}
websocket_connections: List[WebSocket] = []

class ConnectionManager:
    def __init__(self):
        self.active_connections: List[WebSocket] = []

    async def connect(self, websocket: WebSocket):
        await websocket.accept()
        self.active_connections.append(websocket)

    def disconnect(self, websocket: WebSocket):
        self.active_connections.remove(websocket)

    async def send_personal_message(self, message: str, websocket: WebSocket):
        await websocket.send_text(message)

    async def broadcast(self, message: str):
        for connection in self.active_connections:
            await connection.send_text(message)

manager = ConnectionManager()

@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})

@app.post("/api/upload/syllabus")
async def upload_syllabus(file: UploadFile = File(...)):
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="教学大纲文件必须是.docx格式")
    
    file_id = str(uuid.uuid4())
    file_path = os.path.join(uploads_dir, f"{file_id}_{file.filename}")
    
    try:
        async with aiofiles.open(file_path, 'wb') as f:
            content = await file.read()
            await f.write(content)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文件保存失败: {str(e)}")
    
    uploaded_files[file_id] = {
        "filename": file.filename,
        "filepath": file_path,
        "type": "syllabus",
        "status": "uploaded"
    }
    
    return {"file_id": file_id, "filename": file.filename, "status": "success"}

@app.post("/api/upload/schedule")
async def upload_schedule(file: UploadFile = File(...)):
    if not file.filename.endswith(('.xlsx', '.xls')):
        raise HTTPException(status_code=400, detail="教学计划文件必须是Excel格式")
    
    file_id = str(uuid.uuid4())
    file_path = os.path.join(uploads_dir, f"{file_id}_{file.filename}")
    
    try:
        async with aiofiles.open(file_path, 'wb') as f:
            content = await file.read()
            await f.write(content)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文件保存失败: {str(e)}")
    
    uploaded_files[file_id] = {
        "filename": file.filename,
        "filepath": file_path,
        "type": "schedule",
        "status": "uploaded"
    }
    
    return {"file_id": file_id, "filename": file.filename, "status": "success"}

@app.post("/api/upload/template")
async def upload_template(file: UploadFile = File(...)):
    if not file.filename.endswith('.docx'):
        raise HTTPException(status_code=400, detail="教案模板文件必须是.docx格式")
    
    file_id = str(uuid.uuid4())
    file_path = os.path.join(uploads_dir, f"{file_id}_{file.filename}")
    
    try:
        async with aiofiles.open(file_path, 'wb') as f:
            content = await file.read()
            await f.write(content)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文件保存失败: {str(e)}")
    
    uploaded_files[file_id] = {
        "filename": file.filename,
        "filepath": file_path,
        "type": "template",
        "status": "uploaded"
    }
    
    return {"file_id": file_id, "filename": file.filename, "status": "success"}

@app.get("/api/files")
async def get_files():
    return {"files": uploaded_files}

@app.delete("/api/files/{file_id}")
async def delete_file(file_id: str):
    if file_id not in uploaded_files:
        raise HTTPException(status_code=404, detail="文件不存在")
    
    file_info = uploaded_files[file_id]
    if os.path.exists(file_info["filepath"]):
        os.remove(file_info["filepath"])
    
    del uploaded_files[file_id]
    return {"status": "success"}

from fastapi.exceptions import RequestValidationError
from fastapi.responses import JSONResponse

@app.exception_handler(RequestValidationError)
async def validation_exception_handler(request, exc):
    print(f"=== 请求验证错误 ===")
    print(f"请求URL: {request.url}")
    print(f"请求方法: {request.method}")
    print(f"请求头: {request.headers}")
    print(f"错误详情: {exc}")
    return JSONResponse(
        status_code=422,
        content={"detail": exc.errors(), "body": exc.body}
    )

@app.post("/api/parse/{file_type}")
async def parse_file(file_type: str, request: ParseRequest):
    # 详细的请求日志记录
    timestamp = time.strftime("%Y-%m-%d %H:%M:%S")
    print(f"\n=== [{timestamp}] 解析API调用 ===")
    print(f"文件类型: {file_type}")
    print(f"请求对象: {request}")
    print(f"请求类型: {type(request)}")
    print(f"请求file_id: {request.file_id}")
    
    # 参数验证
    file_id = request.file_id
    if not file_id:
        error_msg = "缺少文件ID"
        print(f"✗ 参数验证失败: {error_msg}")
        raise HTTPException(status_code=400, detail=error_msg)
    
    if file_id not in uploaded_files:
        error_msg = f"文件ID {file_id} 不存在"
        print(f"✗ 文件验证失败: {error_msg}")
        print(f"当前已上传文件: {list(uploaded_files.keys())}")
        raise HTTPException(status_code=404, detail=error_msg)
    
    file_info = uploaded_files[file_id]
    if file_info["type"] != file_type:
        error_msg = f"文件类型错误: 期望 {file_type}, 实际 {file_info['type']}"
        print(f"✗ 文件类型验证失败: {error_msg}")
        raise HTTPException(status_code=400, detail=error_msg)
    
    # 详细的文件信息日志
    file_path = file_info["filepath"]
    print(f"解析请求详情:")
    print(f"  - 文件ID: {file_id}")
    print(f"  - 文件类型: {file_type}")
    print(f"  - 文件路径: {file_path}")
    print(f"  - 文件存在: {os.path.exists(file_path)}")
    
    if os.path.exists(file_path):
        file_size = os.path.getsize(file_path)
        file_readable = os.access(file_path, os.R_OK)
        print(f"  - 文件大小: {file_size} 字节")
        print(f"  - 文件可读: {file_readable}")
        print(f"  - 文件扩展名: {os.path.splitext(file_path)[1]}")
    else:
        error_msg = f"文件不存在: {file_path}"
        print(f"✗ 文件访问失败: {error_msg}")
        raise HTTPException(status_code=404, detail=error_msg)
    
    print(f"  - 当前工作目录: {os.getcwd()}")
    print(f"  - uploads目录: {uploads_dir}")
    print(f"  - uploads目录存在: {os.path.exists(uploads_dir)}")
    
    try:
        print(f"开始解析 {file_type} 文件...")
        
        if file_type == "schedule":
            # 处理Excel教学进度表
            print("  步骤1: 验证Excel文件结构...")
            validation_result = DataParser.validate_excel_structure(file_info["filepath"])
            
            print(f"  验证结果: {validation_result}")
            if not validation_result["valid"]:
                error_details = "文件结构验证失败:\n" + "\n".join(validation_result["issues"])
                print(f"  ✗ Excel文件格式验证失败")
                print(f"  错误详情: {error_details}")
                return {
                    "status": "error",
                    "detail": "Excel文件格式不正确",
                    "error_details": error_details,
                    "validation_result": validation_result
                }
            
            print("  步骤2: 解析Excel文件内容...")
            schedule_data = DataParser.parse_schedule(file_info["filepath"])
            
            print(f"  ✓ 解析成功，共 {len(schedule_data)} 条记录")
            uploaded_files[file_id]["parsed_data"] = schedule_data
            uploaded_files[file_id]["status"] = "parsed"
            
            return {
                "status": "success", 
                "data": schedule_data,
                "message": f"成功解析{len(schedule_data)}条课程记录"
            }
            
        elif file_type == "syllabus":
            # 处理Word教学大纲
            print("  步骤1: 解析Word教学大纲...")
            syllabus_data = DataParser.parse_syllabus(file_info["filepath"])
            
            word_count = syllabus_data.get('word_count', 0)
            paragraph_count = syllabus_data.get('paragraph_count', 0)
            
            print(f"  ✓ 解析成功，共 {word_count} 字，{paragraph_count} 段")
            uploaded_files[file_id]["parsed_data"] = syllabus_data
            uploaded_files[file_id]["status"] = "parsed"
            
            return {
                "status": "success", 
                "data": syllabus_data,
                "message": f"成功解析教学大纲 (共{word_count}字，{paragraph_count}段)"
            }
            
        elif file_type == "template":
            # 处理Word教案模板
            print("  步骤1: 解析Word教案模板...")
            template_data = DataParser.parse_syllabus(file_info["filepath"])
            
            word_count = template_data.get('word_count', 0)
            paragraph_count = template_data.get('paragraph_count', 0)
            
            print(f"  ✓ 解析成功，共 {word_count} 字，{paragraph_count} 段")
            uploaded_files[file_id]["parsed_data"] = template_data
            uploaded_files[file_id]["status"] = "parsed"
            
            return {
                "status": "success", 
                "data": template_data,
                "message": f"成功解析教案模板 (共{word_count}字，{paragraph_count}段)"
            }
            
        else:
            error_msg = f"不支持的文件类型: {file_type}"
            print(f"  ✗ {error_msg}")
            raise HTTPException(status_code=400, detail=error_msg)
        
    except ValueError as e:
        # 处理已知的验证错误
        error_msg = f"文件解析失败: {str(e)}"
        print(f"  ✗ {error_msg}")
        return {
            "status": "error",
            "detail": "文件解析失败",
            "error_details": str(e)
        }
    except Exception as e:
        # 处理未知错误
        import traceback
        error_traceback = traceback.format_exc()
        print(f"  ✗ 解析{file_type}文件时发生未知错误:")
        print(f"  错误信息: {str(e)}")
        print(f"  错误堆栈: {error_traceback}")
        
        return {
            "status": "error",
            "detail": "文件解析失败",
            "error_details": f"未知错误: {str(e)}"
        }

# 保留原有的具体路由以确保兼容性
@app.post("/api/parse/schedule")
async def parse_schedule(request: ParseRequest):
    return await parse_file("schedule", request)

@app.post("/api/parse/syllabus")
async def parse_syllabus(request: ParseRequest):
    return await parse_file("syllabus", request)

@app.post("/api/parse/template")
async def parse_template(request: ParseRequest):
    return await parse_file("template", request)

@app.post("/api/test/parse")
async def test_parse(request: ParseRequest):
    """测试解析API的请求处理"""
    print(f"=== 测试解析API ===")
    print(f"接收到的请求: {request}")
    print(f"file_id: {request.file_id}")
    return {"status": "test", "file_id": request.file_id}

@app.post("/api/debug/file/{file_id}")
async def debug_file(file_id: str):
    """调试接口：检查文件状态和内容"""
    if file_id not in uploaded_files:
        raise HTTPException(status_code=404, detail="文件不存在")
    
    file_info = uploaded_files[file_id]
    file_path = file_info["filepath"]
    
    debug_info = {
        "file_id": file_id,
        "filename": file_info["filename"],
        "file_type": file_info["type"],
        "status": file_info.get("status", "unknown"),
        "file_path": file_path,
        "file_exists": os.path.exists(file_path)
    }
    
    if debug_info["file_exists"]:
        try:
            debug_info["file_size"] = os.path.getsize(file_path)
            debug_info["file_readable"] = os.access(file_path, os.R_OK)
            
            # 读取文件头部内容进行预览
            if file_info["type"] == "schedule":
                # Excel文件预览
                try:
                    import pandas as pd
                    df = pd.read_excel(file_path, nrows=5)
                    debug_info["preview"] = {
                        "columns": list(df.columns),
                        "rows": len(df),
                        "sample_data": df.head(3).to_dict("records")
                    }
                except Exception as e:
                    debug_info["preview_error"] = str(e)
                    
            elif file_info["type"] in ["syllabus", "template"]:
                # Word文件预览
                try:
                    from docx import Document
                    doc = Document(file_path)
                    paragraphs = [p.text.strip() for p in doc.paragraphs[:5] if p.text.strip()]
                    debug_info["preview"] = {
                        "paragraph_count": len(doc.paragraphs),
                        "sample_paragraphs": paragraphs
                    }
                except Exception as e:
                    debug_info["preview_error"] = str(e)
                    
        except Exception as e:
            debug_info["error"] = str(e)
    
    return {"debug_info": debug_info}

@app.get("/api/debug/dependencies")
async def debug_dependencies():
    """调试接口：检查依赖包状态"""
    dependencies = {}
    
    try:
        import pandas as pd
        dependencies["pandas"] = {
            "installed": True,
            "version": pd.__version__
        }
    except ImportError as e:
        dependencies["pandas"] = {
            "installed": False,
            "error": str(e)
        }
    
    try:
        from docx import Document
        dependencies["python-docx"] = {
            "installed": True,
            "version": getattr(Document, '__version__', 'unknown')
        }
    except ImportError as e:
        dependencies["python-docx"] = {
            "installed": False,
            "error": str(e)
        }
    
    try:
        import openpyxl
        dependencies["openpyxl"] = {
            "installed": True,
            "version": openpyxl.__version__
        }
    except ImportError as e:
        dependencies["openpyxl"] = {
            "installed": False,
            "error": str(e)
        }
    
    return {"dependencies": dependencies}

@app.post("/api/generate")
async def generate_lesson_plans(request: GenerateRequest):
    print(f"=== 生成教案请求 ===")
    print(f"请求对象: {request}")
    print(f"schedule_file_id: {request.schedule_file_id}")
    print(f"syllabus_file_id: {request.syllabus_file_id}")
    print(f"template_file_id: {request.template_file_id}")
    print(f"week_range: {request.week_range}")
    
    # 验证文件存在
    if request.schedule_file_id not in uploaded_files:
        raise HTTPException(status_code=404, detail="教学计划文件不存在")
    
    # 创建生成任务
    task_id = str(uuid.uuid4())
    generation_tasks[task_id] = {
        "status": "pending",
        "progress": 0,
        "total": 0,
        "current": "",
        "error": None,
        "result_files": []
    }
    
    # 启动后台生成任务
    asyncio.create_task(generate_lesson_plans_background(
        task_id, request.schedule_file_id, request.syllabus_file_id, request.template_file_id, request.week_range
    ))
    
    return {"task_id": task_id, "status": "started"}

async def generate_lesson_plans_background(
    task_id: str,
    schedule_file_id: str,
    syllabus_file_id: str = None,
    template_file_id: str = None,
    week_range: str = None
):
    print(f"=== 后台生成任务开始 ===")
    print(f"task_id: {task_id}")
    print(f"schedule_file_id: {schedule_file_id}")
    print(f"syllabus_file_id: {syllabus_file_id}")
    print(f"template_file_id: {template_file_id}")
    print(f"week_range: {week_range}")
    
    try:
        # 验证任务存在
        if task_id not in generation_tasks:
            print(f"错误: 任务 {task_id} 不存在")
            return
        
        # 更新任务状态
        generation_tasks[task_id]["status"] = "running"
        generation_tasks[task_id]["progress"] = 0
        generation_tasks[task_id]["error"] = None
        print(f"任务状态已更新为running")
        
        # 广播状态更新
        progress_message = json.dumps({
            "type": "progress",
            "task_id": task_id,
            "status": "running",
            "progress": 0,
            "message": "开始生成教案...",
            "current": ""
        })
        print(f"广播进度消息: {progress_message}")
        await manager.broadcast(progress_message)
        
        # 获取文件路径
        schedule_file = uploaded_files[schedule_file_id]["filepath"]
        syllabus_file = uploaded_files[syllabus_file_id]["filepath"] if syllabus_file_id else None
        template_file = uploaded_files[template_file_id]["filepath"] if template_file_id else "test_data/template.docx"
        
        # 解析数据
        await manager.broadcast(json.dumps({
            "type": "progress",
            "task_id": task_id,
            "status": "running",
            "progress": 10,
            "message": "正在解析教学进度表...",
            "current": ""
        }))
        
        schedule_data = DataParser.parse_schedule(schedule_file)
        syllabus_data = DataParser.parse_syllabus(syllabus_file) if syllabus_file else None
        
        # 应用周次范围过滤
        if week_range:
            try:
                start_week, end_week = map(int, week_range.split('-'))
                schedule_data = [lesson for lesson in schedule_data if start_week <= lesson['week'] <= end_week]
            except ValueError:
                pass
        
        # 初始化AI生成器和文档生成器
        await manager.broadcast(json.dumps({
            "type": "progress",
            "task_id": task_id,
            "status": "running",
            "progress": 20,
            "message": "正在初始化AI生成器...",
            "current": ""
        }))
        
        print("正在初始化AI生成器...")
        try:
            ai_generator = AIGenerator()
            print("AI生成器初始化成功")
        except Exception as e:
            print(f"AI生成器初始化失败: {e}")
            raise
        
        try:
            doc_builder = DocumentBuilder(template_file)
            print("文档生成器初始化成功")
        except Exception as e:
            print(f"文档生成器初始化失败: {e}")
            raise
        
        # 设置总进度
        total_lessons = len(schedule_data)
        generation_tasks[task_id]["total"] = total_lessons
        
        # 生成教案
        output_dir = "lesson_plans"
        os.makedirs(output_dir, exist_ok=True)
        
        for i, lesson_data in enumerate(schedule_data):
            # 检查任务是否仍然存在
            if task_id not in generation_tasks:
                print(f"任务 {task_id} 已被删除，停止生成")
                break
            
            # 更新当前进度信息
            current_lesson = f"第{lesson_data['week']}周第{lesson_data['lesson']}次课"
            generation_tasks[task_id]["current"] = current_lesson
            
            progress = 20 + int((i / total_lessons) * 70)
            generation_tasks[task_id]["progress"] = progress
            
            # 广播进度更新
            progress_message = json.dumps({
                "type": "progress",
                "task_id": task_id,
                "status": "running",
                "progress": progress,
                "message": f"正在生成{current_lesson}教案...",
                "current": current_lesson
            })
            print(f"广播进度消息: {progress_message}")
            await manager.broadcast(progress_message)
            
            # 生成AI内容
            ai_content = {}
            for field in ["单元教学目标", "教学重点", "教学难点", "教学活动", "作业布置", "教学资源", "教学反思", "教学评价"]:
                print(f"=== 生成 {field} ===")
                print(f"lesson_data: {lesson_data}")
                print(f"lesson_data keys: {lesson_data.keys() if isinstance(lesson_data, dict) else 'Not a dict'}")
                ai_content[field] = ai_generator.generate_content(field, lesson_data=lesson_data, syllabus_data=syllabus_data)
            
            # 生成文档
            output_filename = f"第{lesson_data['week']}周第{lesson_data['lesson']}次课教案.docx"
            output_path = os.path.join(output_dir, output_filename)
            
            doc_builder.build_lesson_plan(lesson_data, ai_content, output_path)
            generation_tasks[task_id]["result_files"].append(output_filename)
            
            print(f"已生成: {output_filename}")
        
        # 完成生成
        generation_tasks[task_id]["status"] = "completed"
        generation_tasks[task_id]["progress"] = 100
        
        completion_message = json.dumps({
            "type": "progress",
            "task_id": task_id,
            "status": "completed",
            "progress": 100,
            "message": f"教案生成完成！共生成{len(generation_tasks[task_id]['result_files'])}个教案",
            "current": ""
        })
        print(f"广播完成消息: {completion_message}")
        await manager.broadcast(completion_message)
        
    except Exception as e:
        print(f"生成任务异常: {str(e)}")
        import traceback
        print(f"错误堆栈: {traceback.format_exc()}")
        
        generation_tasks[task_id]["status"] = "failed"
        generation_tasks[task_id]["error"] = str(e)
        
        error_message = json.dumps({
            "type": "progress",
            "task_id": task_id,
            "status": "failed",
            "error": str(e),
            "message": f"生成失败: {str(e)}",
            "current": ""
        })
        print(f"广播错误消息: {error_message}")
        await manager.broadcast(error_message)

@app.get("/api/generate/status/{task_id}")
async def get_generation_status(task_id: str):
    if task_id not in generation_tasks:
        raise HTTPException(status_code=404, detail="任务不存在")
    
    return generation_tasks[task_id]

@app.post("/api/generate/{task_id}/pause")
async def pause_generation(task_id: str):
    print(f"=== 暂停生成任务请求 ===")
    print(f"任务ID: {task_id}")
    print(f"当前活动任务: {list(generation_tasks.keys())}")
    
    # 验证任务ID
    if not task_id:
        print("错误: 任务ID为空")
        raise HTTPException(status_code=400, detail="任务ID不能为空")
    
    if task_id not in generation_tasks:
        print(f"错误: 任务 {task_id} 不存在")
        raise HTTPException(status_code=404, detail="任务不存在")
    
    # 获取任务信息
    task = generation_tasks[task_id]
    print(f"任务当前状态: {task.get('status', 'unknown')}")
    
    # 检查任务是否已经结束
    if task.get("status") in ["completed", "failed", "stopped"]:
        print(f"任务已经处于结束状态: {task.get('status')}")
        return {
            "status": task.get("status"),
            "message": f"任务已经{task.get('status')}"
        }
    
    # 检查任务是否已经暂停
    if task.get("status") == "paused":
        print(f"任务已经处于暂停状态")
        return {
            "status": "paused",
            "message": "任务已经处于暂停状态"
        }
    
    # 检查任务是否可以暂停
    if task.get("status") != "running":
        print(f"任务状态无法暂停: {task.get('status')}")
        raise HTTPException(
            status_code=400, 
            detail=f"任务状态为 {task.get('status')}，无法暂停"
        )
    
    try:
        # 更新任务状态
        task["status"] = "paused"
        print(f"任务 {task_id} 已标记为暂停")
        
        # 广播暂停消息
        pause_message = json.dumps({
            "type": "progress",
            "task_id": task_id,
            "status": "paused",
            "progress": task.get("progress", 0),
            "message": "任务已暂停",
            "current": task.get("current", "")
        })
        
        print(f"广播暂停消息: {pause_message}")
        await manager.broadcast(pause_message)
        
        print(f"任务 {task_id} 暂停成功")
        return {
            "status": "paused",
            "message": "任务已成功暂停",
            "task_id": task_id
        }
        
    except Exception as e:
        print(f"暂停任务时发生错误: {str(e)}")
        import traceback
        print(f"错误堆栈: {traceback.format_exc()}")
        
        # 更新任务状态为失败
        task["status"] = "failed"
        task["error"] = f"暂停失败: {str(e)}"
        
        raise HTTPException(
            status_code=500,
            detail=f"暂停任务失败: {str(e)}"
        )

@app.post("/api/generate/{task_id}/stop")
async def stop_generation(task_id: str):
    print(f"=== 终止生成任务请求 ===")
    print(f"任务ID: {task_id}")
    print(f"当前活动任务: {list(generation_tasks.keys())}")
    
    # 验证任务ID
    if not task_id:
        print("错误: 任务ID为空")
        raise HTTPException(status_code=400, detail="任务ID不能为空")
    
    if task_id not in generation_tasks:
        print(f"错误: 任务 {task_id} 不存在")
        raise HTTPException(status_code=404, detail="任务不存在")
    
    # 获取任务信息
    task = generation_tasks[task_id]
    print(f"任务当前状态: {task.get('status', 'unknown')}")
    
    # 检查任务是否已经结束
    if task.get("status") in ["completed", "failed"]:
        print(f"任务已经处于结束状态: {task.get('status')}")
        return {
            "status": task.get("status"),
            "message": f"任务已经{task.get('status')}"
        }
    
    # 检查任务是否已经终止
    if task.get("status") == "stopped":
        print(f"任务已经处于终止状态")
        return {
            "status": "stopped",
            "message": "任务已经处于终止状态"
        }
    
    try:
        # 更新任务状态
        task["status"] = "stopped"
        task["error"] = "用户终止"
        
        print(f"任务 {task_id} 已标记为停止")
        
        # 广播停止消息
        stop_message = json.dumps({
            "type": "progress",
            "task_id": task_id,
            "status": "stopped",
            "progress": task.get("progress", 0),
            "message": "任务已终止",
            "current": task.get("current", "")
        })
        
        print(f"广播停止消息: {stop_message}")
        await manager.broadcast(stop_message)
        
        print(f"任务 {task_id} 终止成功")
        return {
            "status": "stopped",
            "message": "任务已成功终止",
            "task_id": task_id,
            "progress": task.get("progress", 0),
            "current": task.get("current", "")
        }
        
    except Exception as e:
        print(f"终止任务时发生错误: {str(e)}")
        import traceback
        print(f"错误堆栈: {traceback.format_exc()}")
        
        # 更新任务状态为失败
        task["status"] = "failed"
        task["error"] = f"终止失败: {str(e)}"
        
        raise HTTPException(
            status_code=500,
            detail=f"终止任务失败: {str(e)}"
        )

@app.post("/api/generate/{task_id}/resume")
async def resume_generation(task_id: str):
    print(f"=== 恢复生成任务请求 ===")
    print(f"任务ID: {task_id}")
    print(f"当前活动任务: {list(generation_tasks.keys())}")
    
    # 验证任务ID
    if not task_id:
        print("错误: 任务ID为空")
        raise HTTPException(status_code=400, detail="任务ID不能为空")
    
    if task_id not in generation_tasks:
        print(f"错误: 任务 {task_id} 不存在")
        raise HTTPException(status_code=404, detail="任务不存在")
    
    # 获取任务信息
    task = generation_tasks[task_id]
    print(f"任务当前状态: {task.get('status', 'unknown')}")
    
    # 检查任务是否已经结束
    if task.get("status") in ["completed", "failed", "stopped"]:
        print(f"任务已经处于结束状态: {task.get('status')}")
        return {
            "status": task.get("status"),
            "message": f"任务已经{task.get('status')}"
        }
    
    # 检查任务是否在暂停状态
    if task.get("status") != "paused":
        print(f"任务状态无法恢复: {task.get('status')}")
        raise HTTPException(
            status_code=400, 
            detail=f"任务状态为 {task.get('status')}，无法恢复"
        )
    
    try:
        # 更新任务状态
        task["status"] = "running"
        print(f"任务 {task_id} 已标记为运行")
        
        # 广播恢复消息
        resume_message = json.dumps({
            "type": "progress",
            "task_id": task_id,
            "status": "running",
            "progress": task.get("progress", 0),
            "message": "任务已恢复",
            "current": task.get("current", "")
        })
        
        print(f"广播恢复消息: {resume_message}")
        await manager.broadcast(resume_message)
        
        print(f"任务 {task_id} 恢复成功")
        return {
            "status": "running",
            "message": "任务已成功恢复",
            "task_id": task_id,
            "progress": task.get("progress", 0),
            "current": task.get("current", "")
        }
        
    except Exception as e:
        print(f"恢复任务时发生错误: {str(e)}")
        import traceback
        print(f"错误堆栈: {traceback.format_exc()}")
        
        # 更新任务状态为失败
        task["status"] = "failed"
        task["error"] = f"恢复失败: {str(e)}"
        
        raise HTTPException(
            status_code=500,
            detail=f"恢复任务失败: {str(e)}"
        )

async def resume_generation_background(task_id: str):
    """恢复生成的后台任务"""
    try:
        task = generation_tasks[task_id]
        
        await manager.broadcast(json.dumps({
            "type": "progress",
            "task_id": task_id,
            "status": "running",
            "progress": task["progress"],
            "message": "继续生成教案..."
        }))
        
        # 这里需要从上次暂停的地方继续生成
        # 简化实现：重新开始生成（实际应该记录进度并从断点继续）
        # TODO: 实现真正的断点续传功能
        
    except Exception as e:
        generation_tasks[task_id]["status"] = "failed"
        generation_tasks[task_id]["error"] = str(e)
        
        await manager.broadcast(json.dumps({
            "type": "progress",
            "task_id": task_id,
            "status": "failed",
            "error": str(e),
            "message": f"生成失败: {str(e)}"
        }))

@app.get("/api/generate/results")
async def get_generation_results():
    results = []
    lesson_plans_dir = "lesson_plans"
    
    if os.path.exists(lesson_plans_dir):
        for filename in os.listdir(lesson_plans_dir):
            if filename.endswith('.docx'):
                filepath = os.path.join(lesson_plans_dir, filename)
                stat = os.stat(filepath)
                results.append({
                    "filename": filename,
                    "filepath": filepath,
                    "size": stat.st_size,
                    "created": stat.st_ctime
                })
    
    return {"results": results}

@app.get("/api/download/{filename}")
async def download_file(filename: str):
    lesson_plans_dir = "lesson_plans"
    filepath = os.path.join(lesson_plans_dir, filename)
    
    if not os.path.exists(filepath):
        raise HTTPException(status_code=404, detail="文件不存在")
    
    from fastapi.responses import FileResponse
    return FileResponse(
        path=filepath,
        filename=filename,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.websocket("/ws/progress")
async def websocket_endpoint(websocket: WebSocket):
    await manager.connect(websocket)
    try:
        while True:
            data = await websocket.receive_text()
            # 这里可以处理客户端发送的消息
    except WebSocketDisconnect:
        manager.disconnect(websocket)

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)