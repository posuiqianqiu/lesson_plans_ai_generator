from docx import Document
import os

# Ensure the data directory exists
if not os.path.exists('data'):
    os.makedirs('data')

# --- Create syllabus.docx ---
syllabus_doc = Document()
syllabus_doc.add_heading('课程教学大纲', level=1)
syllabus_doc.add_paragraph('课程名称：Python程序设计')
syllabus_doc.add_paragraph('适用专业：计算机应用技术')
syllabus_doc.add_paragraph('课程目标：掌握Python基础语法，能够编写简单的应用程序。')
syllabus_doc.add_paragraph('教学活动：包括课堂讲授、案例分析、编程练习、小组讨论等。')
syllabus_doc.add_paragraph('教学资源：教材、在线教程、编程工具、实验环境等。')
syllabus_doc.add_paragraph('教学反思：定期回顾教学过程，优化教学方法。')
syllabus_doc.add_paragraph('教学评价：通过作业、测验、项目等多种方式评价学生学习效果。')
syllabus_doc.save('data/syllabus.docx')


# --- Create template.docx ---
template_doc = Document()
template_doc.add_heading('{{course_name}} 教案', level=1)
template_doc.add_paragraph('周次：{{week}}')
template_doc.add_paragraph('课次：{{lesson}}')
template_doc.add_paragraph('章节内容：{{chapter_content}}')
template_doc.add_paragraph('课时：{{class_hours}}')

template_doc.add_heading('教学目标', level=2)
template_doc.add_paragraph('{{教学目标}}')

template_doc.add_heading('教学重点', level=2)
template_doc.add_paragraph('{{教学重点}}')

template_doc.add_heading('教学难点', level=2)
template_doc.add_paragraph('{{教学难点}}')

template_doc.add_heading('教学活动', level=2)
template_doc.add_paragraph('{{教学活动}}')

template_doc.add_heading('教学资源', level=2)
template_doc.add_paragraph('{{教学资源}}')

template_doc.add_heading('教学反思', level=2)
template_doc.add_paragraph('{{教学反思}}')

template_doc.add_heading('教学评价', level=2)
template_doc.add_paragraph('{{教学评价}}')

template_doc.save('data/template.docx')
print("data/template.docx created successfully.")